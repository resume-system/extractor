from docx import Document
from docx.shared import Inches
import datetime
import os
import cv2
import shutil
import numpy as np
import pandas as pd
import re
import json
from tqdm import tqdm

# 设置实体抽取信息
schema = ['姓名', '出生年月', '电话', '性别', '项目名称', '项目责任', '项目时间', '籍贯', '政治面貌', '落户市县', '毕业院校', '学位', '毕业时间', '工作时间', '工作内容', '职务', '工作单位']

def get_paragraphs_text(path):
    document = Document(path) 
    # 有的简历是表格式样的，因此，不仅需要提取正文，还要提取表格
    col_keys = [] # 获取列名
    col_values = [] # 获取列值
    index_num = 0
    # 表格提取中，需要添加一个去重机制
    fore_str = ""
    cell_text = ""
    for table in document.tables:
        for row_index,row in enumerate(table.rows):
            for col_index,cell in enumerate(row.cells):
                if fore_str != cell.text:
                    if index_num % 2==0:
                        col_keys.append(cell.text)
                    else:
                        col_values.append(cell.text)
                    fore_str = cell.text
                    index_num +=1
                    # 避免使用换行符
                    cell_text += cell.text + '；'
    # 提取正文文本
    paragraphs_text = ""
    for paragraph in document.paragraphs:
        # 拼接一个list,包括段落的结构和内容，避免使用换行符
        paragraphs_text += paragraph.text + "；"
    # 剔除掉返回内容中多余的空格、tab、换行符
    cell_text = cell_text.replace('\n', '；').replace(' ', '').replace('\t', '')
    paragraphs_text = paragraphs_text.replace('\n', '；').replace(' ', '').replace('\t', '')
    return cell_text, paragraphs_text

with open('/resume_train/train.json', "r", encoding="utf-8") as f1, open('unlabeled_data.txt', "a", encoding='utf-8') as f2:
    raw_examples = json.loads(f1.read())
    label_list = []    
    line_num = 1
    for line in  tqdm(raw_examples): 
        result_list = []
        # 解析简历文本内容
        cell_text, paragraphs_text = get_paragraphs_text(os.path.join('resume_train_20200121/docx',line) + '.docx')
        text_content = cell_text + paragraphs_text
        # 保存提取的简历内容到无标签数据文件
        f2.write(text_content + '\n')       
        for item in schema:
            schema_dict = {} 
            if item in raw_examples[line] and text_content.find(raw_examples[line][item]) > 0:
                # 找到要抽取的文本内容
                schema_dict["text"] = raw_examples[line][item]
                # 遍历字符串，找到首个符合匹配的字符位置
                schema_dict["start"] = text_content.find(raw_examples[line][item])
                # 计算文本内容结束位置
                schema_dict["end"] = len(raw_examples[line][item]) + text_content.find(raw_examples[line][item])
                # 保存标签信息
                schema_dict["labels"] = [item]
            if '项目经历' in raw_examples[line]:
                    for i in range(len(raw_examples[line]['项目经历'])):
                        if item in raw_examples[line]['项目经历'][i] and text_content.find(raw_examples[line]['项目经历'][i][item]) > 0:
                            schema_dict["text"] = raw_examples[line]['项目经历'][i][item]
                            schema_dict["start"] = text_content.find(raw_examples[line]['项目经历'][i][item])
                            schema_dict["end"] = len(raw_examples[line]['项目经历'][i][item]) + text_content.find(raw_examples[line]['项目经历'][i][item])
                            schema_dict["labels"] = [item]
            if '工作经历' in raw_examples[line]:
                for i in range(len(raw_examples[line]['工作经历'])):
                    if item in raw_examples[line]['工作经历'][i] and text_content.find(raw_examples[line]['工作经历'][i][item]) > 0:
                        schema_dict["text"] = raw_examples[line]['工作经历'][i][item]
                        schema_dict["start"] = text_content.find(raw_examples[line]['工作经历'][i][item])
                        schema_dict["end"] = len(raw_examples[line]['工作经历'][i][item]) + text_content.find(raw_examples[line]['工作经历'][i][item]) - 1
                        schema_dict["labels"] = [item]                                             
            if '教育经历' in raw_examples[line]:
                for i in range(len(raw_examples[line]['教育经历'])):
                    if item in raw_examples[line]['教育经历'][i] and text_content.find(raw_examples[line]['教育经历'][i][item]) > 0:
                        schema_dict["text"] = raw_examples[line]['教育经历'][i][item]
                        schema_dict["start"] = text_content.find(raw_examples[line]['教育经历'][i][item])
                        schema_dict["end"] = len(raw_examples[line]['教育经历'][i][item]) + text_content.find(raw_examples[line]['教育经历'][i][item])
                        schema_dict["labels"] = [item]
            if len(schema_dict) > 0:
                result_dict = {"value":schema_dict,
                "id": "",
                "from_name": "label",
                "to_name": "text",
                "type": "labels",
                "origin": "manual"}
                result_list.append(result_dict)
        line_dict = {"id": line_num,
        "annotations":[{"id":line_num,"result":result_list}],
        "data": {"text":text_content}
        }
        label_list.append(line_dict)
        line_num += 1

json.dump(label_list, open('label_studio.json', mode='w'), ensure_ascii=False, indent=4) 